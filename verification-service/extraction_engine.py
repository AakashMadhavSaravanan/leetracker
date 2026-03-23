import io
import pdfplumber
import docx

def extract_pdf(file_bytes: bytes) -> dict:
    try:
        content_text = ""
        page_count = 0
        word_count = 0
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content_text += text + "\n"
                    
        if not content_text.strip():
            return {
                "success": True,
                "content_text": "",
                "word_count": 0,
                "page_count": page_count,
                "error": "No extractable text found (Image-only PDF)" # Stated as a warning in payload technically but keeping it as a text string
            }
            
        word_count = len(content_text.split())
        return {
            "success": True,
            "content_text": content_text.strip(),
            "word_count": word_count,
            "page_count": page_count,
            "error": None
        }
    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypt" in error_msg:
            return {"success": False, "content_text": "", "word_count": 0, "page_count": 0, "error": "File is password protected"}
        
        return {"success": False, "content_text": "", "word_count": 0, "page_count": 0, "error": "Could not read file"}


def extract_docx(file_bytes: bytes) -> dict:
    try:
        content_text = ""
        word_count = 0
        
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_text += para.text + "\n"
                
        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content_text += cell.text + " "
                content_text += "\n"
                
        if not content_text.strip():
             return {
                "success": True,
                "content_text": "",
                "word_count": 0,
                "page_count": 0,
                "error": "No extractable text found"
             }

        word_count = len(content_text.split())
        return {
            "success": True,
            "content_text": content_text.strip(),
            "word_count": word_count,
            "page_count": 0,  # docx doesn't easily yield page count natively without rendering
            "error": None
        }
    except Exception as e:
        return {"success": False, "content_text": "", "word_count": 0, "page_count": 0, "error": "Could not read file"}
